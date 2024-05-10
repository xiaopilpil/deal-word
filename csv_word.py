from docx import Document
from read_csv import read_csv

# 超参数
csv_filename = r'mould/test.csv'  # 输入路径模版word文件名
mould_filename = r'mould/test.docx'  # 输入路径模版csv文件名
result_path = r'result/'  # 输出结果路径


# 打开文档
replace_line = read_csv(csv_filename)
# 遍历每个人
for i in range(1, len(replace_line)):
    doc = Document(mould_filename)
    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        # 遍历每个列表
        for j in range(1, len(replace_line[0])):
            # 检查段落中是否包含要替换的内容
            if r'{{{{{}}}}}'.format(replace_line[0][j]) in paragraph.text:
                # 遍历段落中的每个运行
                flag = 0
                for _ in range(1, len(paragraph.runs)):
                    if paragraph.runs[flag].underline == paragraph.runs[flag + 1].underline:
                        paragraph.runs[flag].text += paragraph.runs[flag + 1].text
                        paragraph._element.remove(paragraph.runs[flag + 1]._element)
                    else:
                        flag += 1
                for run in paragraph.runs:
                    if r'{{{{{}}}}}'.format(replace_line[0][j]) in run.text:
                        # 替换文本并保留字体格式
                        run.text = run.text.replace(
                            r'{{{{{}}}}}'.format(replace_line[0][j]),
                            replace_line[i][j])

    # 遍历文档中的每个表格
    for table in doc.tables:
        # 遍历表格中的每一行
        for row in table.rows:
            # 遍历每一行的单元格
            for cell in row.cells:
                # 遍历每个列表
                for j in range(1, len(replace_line[0])):
                    # 检查单元格中是否包含要替换的内容
                    if r'{{{{{}}}}}'.format(replace_line[0][j]) in cell.text:
                        # 遍历单元格中的每个段落
                        for paragraph in cell.paragraphs:
                            # 遍历段落中的每个运行
                            flag = 0
                            for _ in range(1, len(paragraph.runs)):
                                if paragraph.runs[flag].underline == paragraph.runs[flag + 1].underline:
                                    paragraph.runs[flag].text += paragraph.runs[flag + 1].text
                                    paragraph._element.remove(paragraph.runs[flag + 1]._element)
                                else:
                                    flag += 1
                            for run in paragraph.runs:
                                if r'{{{{{}}}}}'.format(replace_line[0][j]) in run.text:
                                    # 替换文本并保留字体格式
                                    run.text = run.text.replace(
                                        r'{{{{{}}}}}'.format(replace_line[0][j]),
                                        replace_line[i][j])

    # 保存文档，根据姓名保存为不同的文件
    doc.save('{}{}.docx'.format(result_path, replace_line[i][0]))
